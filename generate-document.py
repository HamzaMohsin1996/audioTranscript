from docx import Document

# Create a new Document
doc = Document()

# Add the student's name and matriculation number
doc.add_heading('Hamza Mohsin', level=1)
doc.add_heading('Matriculation Number: 00150450', level=2)

# Add Assignment Title
doc.add_heading('Assignment 2 – Grounded Theory', level=1)

# Section 1: Apply Open Coding and Axial Coding to Your Transcripts
doc.add_heading('1. Apply Open Coding and Axial Coding to Your Transcripts', level=2)

# Open Coding
doc.add_heading('Open Coding', level=3)
doc.add_paragraph('Codes Identified:')
codes = [
    'Frustration', 'Coping Strategies', 'Learning New Tool', 'Support Resources', 'Work Environment',
    'Collaboration', 'Future Use', 'Dedicated Workspace', 'Shared Living Space', 'Working from Home',
    'Online Resources', 'Keyboard Shortcuts', 'Focus Strategies', 'Learning Strategies'
]
for code in codes:
    doc.add_paragraph(code, style='List Bullet')

# Axial Coding
doc.add_heading('Axial Coding', level=3)
categories = {
    "Theme 1: Challenges of Learning a New Technology": [
        "Dealing with frustration", "Utilizing keyboard shortcuts (Technological Use)",
        "Dealing with frustration (User Behavior)", "Finding online resources (Technological Use)",
        "Challenges Faced"
    ],
    "Theme 2: Strategies for Effective Remote Learning": [
        "Developing learning strategies (User Behavior)", "Maintaining focus (User Behavior)"
    ],
    "Theme 3: Impact of Work Environment on Learning": [
        "Importance of a dedicated workspace (User Environment)", "Utilizing a shared living space (User Environment)",
        "Working from home (User Environment)", "Environment Impact"
    ],
    "Theme 4: Importance of Collaboration": [
        "Collaboration with colleagues (User Behavior)"
    ],
    "Theme 5: Future Application": [
        "Future Use"
    ]
}
for theme, subcodes in categories.items():
    doc.add_heading(theme, level=4)
    for subcode in subcodes:
        doc.add_paragraph(subcode, style='List Bullet')

# Section 2: Summarize the Most Important Categories and Reflect on Their Impact on Design Decisions
doc.add_heading('2. Summarize the Most Important Categories and Reflect on Their Impact on Design Decisions', level=2)

summaries = [
    ("Theme 1: Challenges of Learning a New Technology", 
     "Summary: Learning Protopie was challenging due to its newness and the unfamiliar terminology.",
     "Impact on Design: Simplify user interfaces and provide more intuitive onboarding experiences to reduce frustration."),
    ("Theme 2: Strategies for Effective Remote Learning", 
     "Summary: Participants used noise-cancelling headphones, online resources, and forums to cope with learning challenges.",
     "Impact on Design: Incorporate built-in tutorials and support features in the design to facilitate effective learning."),
    ("Theme 3: Impact of Work Environment on Learning", 
     "Summary: Noise and lack of dedicated workspace significantly impacted concentration.",
     "Impact on Design: Design tools that require less concentration or are more resilient to environmental distractions."),
    ("Theme 4: Importance of Collaboration", 
     "Summary: Support from colleagues and online communities was crucial for learning.",
     "Impact on Design: Integrate collaborative features and communication tools to enhance teamwork."),
    ("Theme 5: Future Application", 
     "Summary: Participants see value in applying Protopie to future projects.",
     "Impact on Design: Ensure the tools are versatile and can be integrated into various projects.")
]

for theme, summary, impact in summaries:
    doc.add_heading(theme, level=3)
    doc.add_paragraph(summary)
    doc.add_paragraph(impact)

# Section 3: List the Open Questions and New Hypotheses
doc.add_heading('3. List the Open Questions and New Hypotheses', level=2)

# Open Questions
doc.add_heading('Open Questions', level=3)
questions = [
    "How do different work environments specifically impact the learning process for new tools?",
    "What specific features in Protopie contribute most to user frustration?",
    "How effective are different coping strategies in mitigating frustration?"
]
for question in questions:
    doc.add_paragraph(question, style='List Number')

# New Hypotheses
doc.add_heading('New Hypotheses', level=3)
hypotheses = [
    "Participants with a dedicated workspace will report less frustration and higher productivity.",
    "Providing context-sensitive help and tutorials within Protopie will reduce user frustration.",
    "The ability to customize the learning interface will improve user satisfaction and learning outcomes."
]
for hypothesis in hypotheses:
    doc.add_paragraph(hypothesis, style='List Number')

# Section 4: Modify Your Interview Guideline for the Next Set of Interviews
doc.add_heading('4. Modify Your Interview Guideline for the Next Set of Interviews', level=2)

# Modified Interview Questions
doc.add_heading('Modified Interview Questions', level=3)
interview_questions = [
    "Can you describe your work environment in more detail?",
    "How does your work environment affect your ability to concentrate and learn new tools?",
    "Which specific features of Protopie did you find most challenging?",
    "How do you think these features could be improved?",
    "What specific strategies did you use to cope with frustration?",
    "How effective did you find these strategies?"
]
for question in interview_questions:
    doc.add_paragraph(question, style='List Number')

# Section 5: Conduct One More Interview
doc.add_heading('5. Conduct One More Interview', level=2)

# Summary of New Interview Findings
doc.add_heading('Summary of New Interview Findings:', level=3)

# Open Coding
doc.add_heading('Open Coding:', level=4)
new_open_codes = [
    "Shared Apartment", "Noise and Interruptions", "Noise-Cancelling Headphones", "Animation Features",
    "Interaction Features", "Terminology Confusion", "Intuitive Tutorials", "Interactive Guides", 
    "Clear Labels", "Learning Strategies", "Breaking Down Tasks", "Regular Breaks", "Online Resources", 
    "Dedicated Workspace", "Confidence After Learning"
]
for code in new_open_codes:
    doc.add_paragraph(code, style='List Bullet')

# Axial Coding
doc.add_heading('Axial Coding:', level=4)
new_categories = {
    "Theme 1: Challenges of Learning a New Technology": [
        "Animation Features", "Interaction Features", "Terminology Confusion", "Clear Labels"
    ],
    "Theme 2: Strategies for Effective Remote Learning": [
        "Breaking Down Tasks", "Regular Breaks", "Online Resources", "Learning Strategies"
    ],
    "Theme 3: Impact of Work Environment on Learning": [
        "Shared Apartment", "Noise and Interruptions", "Noise-Cancelling Headphones", "Dedicated Workspace"
    ],
    "Theme 4: Importance of Collaboration": [
        "(No new insights specifically related to collaboration in this interview)"
    ],
    "Theme 5: Future Application": [
        "Confidence After Learning"
    ]
}
for theme, subcodes in new_categories.items():
    doc.add_heading(theme, level=4)
    for subcode in subcodes:
        doc.add_paragraph(subcode, style='List Bullet')

# Section 6: Document the Results
doc.add_heading('6. Document the Results', level=2)

# Results
results = [
    ("Theme 1: Challenges of Learning a New Technology", 
     "Participants found specific features of Protopie challenging, particularly animations and interactions, compounded by confusing terminology."),
    ("Theme 2: Strategies for Effective Remote Learning", 
     "Effective strategies included breaking down tasks, taking regular breaks, and utilizing online resources."),
    ("Theme 3: Impact of Work Environment on Learning", 
     "Shared living spaces and noise interruptions significantly impact concentration and learning effectiveness. Noise-cancelling headphones partially mitigate these issues."),
    ("Theme 5: Future Application", 
     "Despite initial challenges, participants gain confidence through perseverance and effective learning strategies.")
]

for theme, result in results:
    doc.add_heading(theme, level=3)
    doc.add_paragraph(result)

# Section 7: Prepare a Short Presentation for Session of June 18th (5 Slides Max.)
doc.add_heading('7. Prepare a Short Presentation for Session of June 18th (5 Slides Max.)', level=2)

# Slides
slides = [
    ("Slide 1: Title and Introduction", 
     "Title: 'Challenges and Strategies in Learning New Tools: A Grounded Theory Approach'\nBrief introduction to the research topic and objectives."),
    ("Slide 2: Methodology", 
     "Overview of open and axial coding process.\nDescription of participant demographics."),
    ("Slide 3: Key Findings", 
     "Summarize the main categories/themes (Challenges, Strategies, Work Environment).\nKey insights from the data."),
    ("Slide 4: Open Questions and Hypotheses", 
     "List the open questions.\nPresent new hypotheses based on the findings."),
    ("Slide 5: Next Steps and Conclusion", 
     "Outline the next steps in the research process.\nConcluding remarks and implications for design.")
]

for slide, content in slides:
    doc.add_heading(slide, level=3)
    doc.add_paragraph(content)

# Section 8: MAXQDA Project Link
doc.add_heading('8. MAXQDA Project Link', level=2)
doc.add_paragraph('You can access the MAXQDA project file from the following link:')
doc.add_paragraph('https://drive.google.com/file/d/1-JioSxwbfHpYlco91IDMOyUwvX1XbOke/view?usp=sharing')

# Save the document
doc.save('Hamza_Mohsin_Assignment.docx')
